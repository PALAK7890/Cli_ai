"""
Unit tests for document loaders (TXT, MD, PDF, DOCX, HTML) and error handling.
"""

import os
import tempfile
from pathlib import Path
import pytest
import docx
import fitz  # PyMuPDF
from knowledge.loaders.router import LoaderRouter, load_file_safely
from knowledge.loaders.txt import TxtLoader
from knowledge.loaders.pdf import PdfLoader
from knowledge.loaders.docx import DocxLoader
from knowledge.loaders.html_loader import HtmlLoader


@pytest.fixture
def temp_workspace() -> tempfile.TemporaryDirectory:
    """Fixture to manage a temporary directory for test files."""
    return tempfile.TemporaryDirectory()


def test_txt_loader(temp_workspace) -> None:
    """Test plain text loader."""
    workspace_path = Path(temp_workspace.name)
    txt_path = workspace_path / "sample.txt"
    content = "Hello, this is a plain text file."
    txt_path.write_text(content, encoding="utf-8")

    loader = TxtLoader()
    docs = loader.load(txt_path)

    assert len(docs) == 1
    assert docs[0].text == content
    assert docs[0].source_path == str(txt_path.resolve())
    assert docs[0].page_number is None
    assert docs[0].metadata["format"] == "txt"
    assert docs[0].document_id is not None


def test_markdown_loader(temp_workspace) -> None:
    """Test markdown loader (uses TxtLoader)."""
    workspace_path = Path(temp_workspace.name)
    md_path = workspace_path / "sample.md"
    content = "# Heading\n\nSome markdown text content."
    md_path.write_text(content, encoding="utf-8")

    router = LoaderRouter()
    loader = router.get_loader(md_path)
    assert isinstance(loader, TxtLoader)

    docs = loader.load(md_path)
    assert len(docs) == 1
    assert docs[0].text == content
    assert docs[0].metadata["format"] == "md"


def test_pdf_loader(temp_workspace) -> None:
    """Test PDF loader with text extraction, page-level mapping, and empty page skipping."""
    workspace_path = Path(temp_workspace.name)
    pdf_path = workspace_path / "sample.pdf"

    # Create a dummy 3-page PDF:
    # Page 1: normal text
    # Page 2: empty text (should be skipped)
    # Page 3: normal text
    doc = fitz.open()
    
    page1 = doc.new_page()
    page1.insert_text((50, 50), "Page 1 Content")
    
    # Page 2 is left blank
    doc.new_page()
    
    page3 = doc.new_page()
    page3.insert_text((50, 50), "Page 3 Content")
    
    doc.save(str(pdf_path))
    doc.close()

    loader = PdfLoader()
    docs = loader.load(pdf_path)

    # Page 2 is skipped, so we should get 2 documents
    assert len(docs) == 2
    
    assert docs[0].text == "Page 1 Content"
    assert docs[0].page_number == 1
    assert docs[0].metadata["total_pages"] == 3
    assert docs[0].metadata["format"] == "pdf"

    assert docs[1].text == "Page 3 Content"
    assert docs[1].page_number == 3
    assert docs[1].metadata["total_pages"] == 3


def test_docx_loader(temp_workspace) -> None:
    """Test DOCX paragraph merging loader."""
    workspace_path = Path(temp_workspace.name)
    docx_path = workspace_path / "sample.docx"

    # Create a simple DOCX file
    doc = docx.Document()
    doc.add_paragraph("Paragraph One")
    doc.add_paragraph("")  # empty paragraph to be skipped
    doc.add_paragraph("Paragraph Two")
    doc.save(str(docx_path))

    loader = DocxLoader()
    docs = loader.load(docx_path)

    assert len(docs) == 1
    assert docs[0].text == "Paragraph One\n\nParagraph Two"
    assert docs[0].page_number is None
    assert docs[0].metadata["format"] == "docx"
    assert docs[0].metadata["paragraph_count"] == 2


def test_html_loader(temp_workspace) -> None:
    """Test HTML text stripping loader."""
    workspace_path = Path(temp_workspace.name)
    html_path = workspace_path / "sample.html"

    html_content = """
    <html>
        <head>
            <style>body { color: red; }</style>
            <script>console.log('hello');</script>
        </head>
        <body>
            <h1>Title</h1>
            <p>Hello HTML Content</p>
        </body>
    </html>
    """
    html_path.write_text(html_content, encoding="utf-8")

    loader = HtmlLoader()
    docs = loader.load(html_path)

    assert len(docs) == 1
    # Style and script elements must be removed.
    assert "body { color: red; }" not in docs[0].text
    assert "console.log" not in docs[0].text
    assert "Title" in docs[0].text
    assert "Hello HTML Content" in docs[0].text
    assert docs[0].metadata["format"] == "html"


def test_corrupted_file_handling(temp_workspace) -> None:
    """Verify that loading a corrupted file does not crash the system."""
    workspace_path = Path(temp_workspace.name)
    corrupt_pdf = workspace_path / "corrupt.pdf"
    
    # Write garbage bytes into a .pdf extension
    corrupt_pdf.write_bytes(b"NOT_A_VALID_PDF_DATA")

    # Safe loader wrapper should handle parsing failures and return empty list
    docs = load_file_safely(corrupt_pdf)
    assert docs == []
