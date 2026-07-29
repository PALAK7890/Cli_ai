"""
Loader for DOCX documents.
"""

from pathlib import Path
from typing import List
import docx
from knowledge.loaders.base import BaseLoader, LoadedDocument, compute_document_id


class DocxLoader(BaseLoader):
    """Parses Word document (.docx) files using python-docx."""

    def load(self, path: Path) -> List[LoadedDocument]:
        """
        Loads a DOCX document as a single LoadedDocument.
        
        Args:
            path: Path to the DOCX document.
            
        Returns:
            A list containing a single LoadedDocument with the entire text.
        """
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # We join paragraphs with double newlines to retain paragraph separation
        full_text = "\n\n".join(paragraphs).strip()
        
        doc_id = compute_document_id(path)
        
        return [
            LoadedDocument(
                text=full_text,
                source_path=str(path.resolve()),
                document_id=doc_id,
                page_number=None,
                metadata={
                    "format": "docx",
                    "paragraph_count": len(paragraphs)
                }
            )
        ]
